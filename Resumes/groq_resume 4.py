#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

import os
import time
import json
from pydantic import BaseModel
from groq import BadRequestError, Groq
import pdfplumber
from docx import Document as docxdocument
from spire.doc import *
from spire.doc.common import *
from zipfile import BadZipFile
import openpyxl

class Education(BaseModel):
    institute_name: str
    education_level: str
    year_of_passing: str
    score: str

class Professional_Experience(BaseModel):
    organization_name: str
    profile_title: str
    experience: float
    profile_skills: list[str]

class Resume_Data(BaseModel):
    name: str
    email: list[str]
    phone: list[str]
    city: str
    total_experience: float
    highest_education: str
    education: list[Education]
    professional_experience: list[Professional_Experience]
    sap_skills: list[str]
    other_skills: list[str]

def Resume_Parser(resume_text: str) -> Resume_Data:
    
    client=Groq(api_key=os.getenv('groq_api_key'),)
    prompt="""
Please extract the following JSON structured data from the unstructured resume text provided. The details of the JSON schema are as below:

- Name: The full name of the candidate.
- Email: A list of emails (if multiple, return all emails).
- Phone: A list of phone numbers (if multiple, return all numbers).
- Total_experience: The candidate's total professional experience, calculated as the sum of all individual professional experiences from different organizations. Note: If a job duration is listed as "present," use 2024 as the current year.
- Education: A list of education details, where each entry should include:
  - Institute_name: Name of school, college, university
  - Education level: Examples maybe High School, Bachelor's, Master's, etc.
  - Year of passing: Passing year
  - Score: Score in cgpa or percentage
- Professional_experience: A list of organization experiences, where each entry should include:
  - organization_name: name of the company
  - profile_title: Latest desination in the organization
  - experience: Total experience in the organization in years as float. Note: If a job duration is listed as "present," use 2024 as the current year.
  - profile_skills: A list of skills related to the profile
  Note: If the resume lists projects instead of explicit job experience within a company, summarize the projects in the format of the professional_experience and avoid listing individual projects.
- sap_skills: A list of all SAP-related skills mentioned throughout the resume, including those from the professional section, projects section, descriptions,etc.
- other_skills: A list of all skills that are not SAP-related mentioned throughout the resume, including those from the professional, projects section, descriptions,etc.

Return the data strictly in the above JSON format."""

    try:
        chat_completion=client.chat.completions.create(
            messages=[
                {
                    "role":"system",
                    "content": f"You are an AI resume parser.{prompt} The output resume data should be in this JSON Schema: {json.dumps(Resume_Data.model_json_schema(),indent=2)}",    
                },
                {
                    "role":"user",
                    "content": resume_text,
                }
            ],
            model="llama-3.1-70b-versatile",
            tool_choice="auto",
            response_format={"type":"json_object"},
        )
    
        return chat_completion.choices[0].message.content
    except Exception:
        return None

# Function to extract text and tables from a PDF file
def extract_from_pdf(file_path):
    output = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Extract text
            page_text = page.extract_text() or ""
            output += page_text + "\n"
            
            # Extract and format tables
            tables = page.extract_tables()
            for table in tables:
                output += "\n[Table Start]\n"
                for row in table:
                    row = [cell if cell is not None else "" for cell in row]
                    output += "\t".join(row) + "\n"  # Join each cell by a tab for better formatting
                output += "[Table End]\n"
    return output

# Function to extract text and tables from a DOCX file
def extract_from_docx(file_path):
    output = ""
    try:
        doc = docxdocument(file_path)
    except Exception:
        print(f"Error: The file {file_path} is corrupted.")
        return None
    
    # Iterate through the document elements in order
    for element in doc.element.body:
        if element.tag.endswith("tbl"):  # Table element
            for table in doc.tables:
                if table._element == element:
                    output += "\n[Table Start]\n"
                    for row in table.rows:
                        output += "\t".join([cell.text for cell in row.cells]) + "\n"
                    output += "[Table End]\n"
                    break
        else:  # Other elements (paragraphs, headers, footers, etc.)
            for para in doc.paragraphs:
                if para._element == element:
                    output += para.text + "\n"
                    break
    
    return output


# Function to convert a DOC file to DOCX
def convert_doc_to_docx(resume_dir):
    for resume_file in os.listdir(resume_dir):
        file_path = resume_dir+"/"+resume_file
        if file_path.endswith('.doc'):
            document = Document()
            document.LoadFromFile(file_path)
            document.SaveToFile(file_path+'x', FileFormat.Docx2019)
            document.Close()

            # Remove the original .doc file
            doc_dir=resume_dir+"/"+"doc files"
            if not os.path.exists(doc_dir):
                os.makedirs(doc_dir)  # Create the output directory if it doesn't exist

            try:
                os.rename(file_path, doc_dir+"/"+resume_file)
            except Exception as e:
                print(f"Error removing {file_path}: {e}")

# Function to handle file extraction based on the file type
def extract_text_and_tables(resume_file):      
    if resume_file.endswith(".pdf"):
        resume_text=extract_from_pdf(resume_file)

    elif resume_file.endswith(".docx"):
        resume_text=extract_from_docx(resume_file)

    else:
        return None    
      
    return resume_text

if __name__=="__main__":
    # Directory to scan for files

    out_count=0

    resume_dir = "C:/User/10743104/Documents/DOCS/Resume/CVs"  # Replace with your directory path

    parsed_dir=resume_dir+"/"+"Parsed Resumes"

    if not os.path.exists(parsed_dir):
        os.makedirs(parsed_dir)  # Create the parsed resumes directory if it doesn't exist

    output_dir=resume_dir+"/"+"Output"

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)  # Create the output directory if it doesn't exist

    excel_file=resume_dir+".xlsx"
    excel_file = list(os.path.split(excel_file))
    excel_file=output_dir+"/"+excel_file[1]   # excel file path

    if not os.path.exists(excel_file):
        wb = openpyxl.Workbook()  # Create the excel if it doesn't exists
        wb.save(excel_file) 

    wb = openpyxl.load_workbook(excel_file)  # Load the existing excel
    ws=wb.active

    headers = ["Sr. No.","File Name", "Resume Score", "Name", "Email", "Phone", "City", "Total Experience", "SAP Skills", "Other Skills", 
           "Organization Name", "Profile Title", "Experience in Years", "Profile Related Skills"]

    ws.append(headers)

    convert_doc_to_docx(resume_dir)

    # Column lists
    general_info_columns = ['A', 'B', 'C', 'D', 'E', 'F', 'G','H', 'I', 'J']
    experience_columns = ['K','L', 'M', 'N']
    row_num=2

    count=1  # Sr. No.
    
    for resume_file in os.listdir(resume_dir):
        file_path = resume_dir+"/"+resume_file
        print(out_count, file_path)
        resume_text=extract_text_and_tables(file_path)
        out_count+=1
        if resume_text:
            resume_out=Resume_Parser(resume_text)
            if resume_out:
                os.rename(file_path,parsed_dir+"/"+resume_file)
            else:
                continue
            resume_json=json.loads(resume_out)

            # Fill general information in columns A to H
            general_info = [count,resume_file,""]

            keys = [
                "name", "email", "phone", "city", "total_experience",
                "sap_skills", "other_skills"
            ]

            for key in keys:
                try:
                    if isinstance(resume_json[key], list):
                        general_info.append(", ".join(map(str, resume_json[key])))
                    else:
                        general_info.append(resume_json[key])
                except KeyError:
                    general_info.append("")

            for col, value in zip(general_info_columns, general_info):
                try:
                    ws[f"{col}{row_num}"] = value
                except ValueError as e:
                    ws[f"{col}{row_num}"] = ""

            # Fill professional experience in columns H to K
            if "professional_experience" in resume_json:
                for exp in resume_json["professional_experience"]:
    
                    exp_info = []

                    keys = [
                        "organization_name", "profile_title", "experience", "profile_skills"
                    ]

                    for key in keys:
                        try:
                            if isinstance(exp[key], list):
                                exp_info.append(", ".join(map(str, exp[key])))
                            else:
                                exp_info.append(exp[key])
                        except KeyError:
                            exp_info.append("")
                    
                    for col, value in zip(experience_columns, exp_info):
                        try:
                            ws[f"{col}{row_num}"] = value
                        except ValueError as e:
                            ws[f"{col}{row_num}"] = ""
                    row_num += 1

            row_num+=1
            count+=1

        # Save the workbook
        wb.save(excel_file)

    print(f"Data has been successfully written to {excel_file}")


